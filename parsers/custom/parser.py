"""
Custom Documents Parser - PRODUCTION VERSION
==============================================
Parse custom documents (Excel, CSV, Word, PDF, Text, JSON)
for chatbot indexing

Supported Formats:
- Excel (.xlsx, .xls)
- CSV (.csv)
- Text (.txt, .md)
- Word (.docx)
- PDF (.pdf)
- JSON (.json)
"""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger

from core.models import Process, Component, ProcessType, ComponentType, SystemType


class CustomDocumentParser:
    """
    Parse custom documents for chatbot indexing

    Features:
    - Multi-format support (Excel, CSV, Word, PDF, Text)
    - Automatic format detection
    - Clean text extraction
    - Metadata preservation
    """

    def __init__(self):
        """Initialize custom document parser"""
        self.supported_extensions = {
            '.xlsx', '.xls',  # Excel
            '.csv',  # CSV
            '.txt', '.md', '.rst',  # Text
            '.docx',  # Word
            '.pdf',  # PDF
            '.json',  # JSON
        }

    def parse_directory(self, documents_path: str) -> Dict[str, Any]:
        """
        Parse custom documents directory

        Args:
            documents_path: Path to custom documents directory

        Returns:
            Dict with 'processes' and 'components' lists
        """
        logger.info(f"Parsing custom documents: {documents_path}")

        processes = []
        components = []

        # Find all supported files
        doc_files = self._find_document_files(documents_path)
        logger.info(f"Found {len(doc_files)} custom document files")

        # Parse each file
        for doc_file in doc_files:
            try:
                doc_data = self._parse_document(doc_file)

                if doc_data:
                    # Convert to Process and Component objects
                    process, process_components = self._convert_to_models(doc_data)

                    if process:
                        processes.append(process)
                        components.extend(process_components)

            except Exception as e:
                logger.error(f"Error parsing {doc_file}: {e}")
                continue

        logger.info(f"✓ Parsed {len(processes)} documents, {len(components)} components")

        return {
            "processes": processes,
            "components": components,
            "summary": {
                "total_processes": len(processes),
                "total_components": len(components),
                "source_path": documents_path,
            }
        }

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a single custom document

        Args:
            file_path: Path to document file

        Returns:
            Dict with 'processes' and 'components'
        """
        logger.info(f"Parsing single document: {file_path}")

        try:
            doc_data = self._parse_document(file_path)

            if doc_data:
                process, components = self._convert_to_models(doc_data)

                return {
                    "processes": [process] if process else [],
                    "components": components,
                }
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")

        return {"processes": [], "components": []}

    def _find_document_files(self, base_path: str) -> List[str]:
        """
        Find all supported document files

        Args:
            base_path: Base directory to search

        Returns:
            List of document file paths
        """
        doc_files = []

        for root, dirs, files in os.walk(base_path):
            for file in files:
                if any(file.lower().endswith(ext) for ext in self.supported_extensions):
                    doc_files.append(os.path.join(root, file))

        return sorted(doc_files)

    def _parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document based on its file type

        Args:
            file_path: Path to document file

        Returns:
            Dict with document data
        """
        file_ext = Path(file_path).suffix.lower()
        file_name = Path(file_path).stem

        logger.info(f"Parsing {file_ext} file: {file_name}")

        # Route to appropriate parser
        if file_ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif file_ext == '.csv':
            return self._parse_csv(file_path)
        elif file_ext in ['.txt', '.md', '.rst']:
            return self._parse_text(file_path)
        elif file_ext == '.docx':
            return self._parse_word(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext == '.json':
            return self._parse_json(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return {}

    def _parse_excel(self, file_path: str) -> Dict[str, Any]:
        """Parse Excel file"""
        try:
            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Convert to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)

                sheets_data.append({
                    'sheet_name': sheet_name,
                    'content': sheet_text,
                    'rows': len(df),
                    'columns': list(df.columns),
                })

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'excel',
                'sheets': sheets_data,
                'total_sheets': len(sheets_data),
            }

        except Exception as e:
            logger.error(f"Error parsing Excel file: {e}")
            logger.info("Falling back to text extraction...")
            return self._parse_as_text_fallback(file_path, 'excel')

    def _parse_csv(self, file_path: str) -> Dict[str, Any]:
        """Parse CSV file"""
        try:
            import pandas as pd

            df = pd.read_csv(file_path)

            # Convert to text
            content = df.to_string(index=False)

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'csv',
                'content': content,
                'rows': len(df),
                'columns': list(df.columns),
            }

        except Exception as e:
            logger.error(f"Error parsing CSV file: {e}")
            return self._parse_as_text_fallback(file_path, 'csv')

    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parse text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'text',
                'content': content,
                'lines': content.count('\n') + 1,
            }

        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            return {}

    def _parse_word(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document"""
        try:
            import docx

            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n\n'.join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(' | '.join(row_data))
                tables_text.append('\n'.join(table_data))

            if tables_text:
                content += '\n\nTables:\n' + '\n\n'.join(tables_text)

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'word',
                'content': content,
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
            }

        except Exception as e:
            logger.error(f"Error parsing Word file: {e}")
            logger.info("Falling back to text extraction...")
            return self._parse_as_text_fallback(file_path, 'word')

    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file"""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)

                # Extract text from all pages
                pages = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text.strip():
                        pages.append(f"Page {i+1}:\n{text}")

                content = '\n\n'.join(pages)

                return {
                    'name': Path(file_path).stem,
                    'file_path': file_path,
                    'file_type': 'pdf',
                    'content': content,
                    'pages': len(pdf.pages),
                }

        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            logger.info("Install PyPDF2 with: pip install PyPDF2")
            return self._parse_as_text_fallback(file_path, 'pdf')

    def _parse_json(self, file_path: str) -> Dict[str, Any]:
        """Parse JSON file"""
        try:
            import json

            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Convert to readable text
            content = json.dumps(data, indent=2)

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'json',
                'content': content,
            }

        except Exception as e:
            logger.error(f"Error parsing JSON file: {e}")
            return self._parse_as_text_fallback(file_path, 'json')

    def _parse_as_text_fallback(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Fallback: try to read as plain text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            return {
                'name': Path(file_path).stem,
                'file_path': file_path,
                'file_type': file_type,
                'content': content,
            }
        except:
            return {}

    def _convert_to_models(self, doc_data: Dict[str, Any]) -> tuple:
        """
        Convert parsed document to Process and Component objects

        Args:
            doc_data: Parsed document data

        Returns:
            Tuple of (Process, List[Component])
        """
        doc_name = doc_data.get('name', 'unknown')
        doc_path = doc_data.get('file_path', '')
        file_type = doc_data.get('file_type', 'unknown')

        # Generate unique hash from file path (cross-platform compatible)
        normalized_path = str(Path(doc_path).as_posix()) if doc_path else ''
        file_hash = hashlib.md5(normalized_path.encode()).hexdigest()[:8]

        # Create Process object with unique ID
        process = Process(
            id=f"custom_{doc_name}_{file_hash}",
            name=doc_name,
            system=SystemType.CUSTOM,
            process_type=ProcessType.DOCUMENT,
            file_path=doc_path,
            description=f"Custom {file_type.upper()} Document: {doc_name}",
        )

        components = []

        # Handle Excel files with multiple sheets
        if file_type == 'excel' and 'sheets' in doc_data:
            for idx, sheet_data in enumerate(doc_data['sheets']):
                component = Component(
                    id=f"{process.id}_sheet_{idx}",
                    name=f"{doc_name} - {sheet_data['sheet_name']}",
                    component_type=ComponentType.DOCUMENT,
                    system="custom",
                    file_path=doc_path,
                    process_id=process.id,
                    process_name=process.name,
                    description=f"Excel sheet with {sheet_data['rows']} rows",
                    raw_content=sheet_data['content'],
                )
                components.append(component)

        # Handle other file types (single component)
        else:
            content = doc_data.get('content', '')

            component = Component(
                id=f"{process.id}_main",
                name=doc_name,
                component_type=ComponentType.DOCUMENT,
                system="custom",
                file_path=doc_path,
                process_id=process.id,
                process_name=process.name,
                description=f"{file_type.upper()} document",
                raw_content=content,
            )
            components.append(component)

        # Update process
        process.component_ids = [c.id for c in components]
        process.component_count = len(components)

        return process, components
